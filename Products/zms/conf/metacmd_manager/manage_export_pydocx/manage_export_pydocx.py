#!/usr/bin/python
# -*- coding: utf-8 -*-

# IMPORT GENERAL LIBRARIES
import os
import re
import shutil
import tempfile
import urllib
import json
import requests
from io import BytesIO
import sys
import datetime

# IMPORT ZMS LIBRARIES
from Products.zms import standard
# from Products.zms import rest_api

# X/HTML LIBRARIES
from bs4 import BeautifulSoup, NavigableString
from lxml import etree

# IMPORT DOCX LIBRARIES
import docx
from docx.oxml import OxmlElement, ns, parse_xml
from docx.oxml.ns import nsdecls
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docx.shared import Emu
from docx.shared import Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.section import WD_SECTION_START


# #############################################
# GLOBALS
# #############################################
# ZMS-Context will be set in main function manage_export_pydocx()
zmscontext = None
# DOCX-Document will be set in main function manage_export_pydocx()
doc = None
# Set local path for docx-template
docx_tmpl = open("/home/zope/instance/zms4_gez/neon-entw/Extensions/neon.docx", "rb")


# #############################################
# Helper Functions 1: DOCX-XML
# e.g. `add_page_number(run)` : add page number to text-run (e.g. footer)
# Hint: the docx API does not support the page counter directly. 
# We have to create a custom footer with a page counter.
# #############################################
# Reference: DocumentFormat.OpenXml
# https://learn.microsoft.com/en-us/dotnet/api/overview/openxml/?view=openxml-3.0.1


# XML-Helpers
def create_element(name):
	return OxmlElement(name)

def create_attribute(element, name, value):
	element.set(ns.qn(name), value)


def get_normalized_image_width(w, h, max_w = 460):
	if w:
		if h > w:
			scale =  h>max_w and float(h)/float(max_w) or 1
		else:
			scale =  w>max_w and float(w)/float(max_w) or 1
		w = int(w/scale)
	else:
		w = max_w
	return w

# #############################################

# ADD DATA FIELD: eg PAGE, SAVEDATE
def add_field(paragraph, field_code="PAGE"):
	fldChar1 = create_element('w:fldChar')
	create_attribute(fldChar1, 'w:fldCharType', 'begin')
	instrText = create_element('w:instrText')
	create_attribute(instrText, 'xml:space', 'preserve')
	instrText.text = field_code
	fldChar2 = create_element('w:fldChar')
	create_attribute(fldChar2, 'w:fldCharType', 'end')
	run = paragraph.add_run()
	run._r.append(fldChar1)
	run._r.append(instrText)
	run._r.append(fldChar2)


# BOOKMARK ZMS-ID
def prepend_bookmark(docx_block, bookmark_id):
	bookmark_start = create_element('w:bookmarkStart')
	create_attribute(bookmark_start, 'w:id', bookmark_id)
	create_attribute(bookmark_start, 'w:name', bookmark_id)
	bookmark_end = create_element('w:bookmarkEnd')
	create_attribute(bookmark_end, 'w:id', bookmark_id)
	try:
		docx_block._element.insert(0, bookmark_end)
		docx_block._element.insert(0, bookmark_start)
	except:
		pass

def set_block_as_listitem(docx_block, list_type='ul', level=0):
	# Set list properties to docx-block
	# ul = List Bullet => numId=2
	# ol = List Number => numId=3
	# level: 0-8
	# Hints:
	# 1. ul/ol/li are handled as blocks in add_htmlblock_to_docx.add_list
	# 2. docx-numbering.xml needs suitable list style definitions with numId=2 and 3
	# Example xml code for a list item paragraph:
		# <w:p>
		# 	<w:pPr>
		# 		<w:pStyle w:val="Normal"/>
		# 		<w:numPr>
		# 			<w:ilvl w:val="0"/>
		# 			<w:numId w:val="2"/>
		# 		</w:numPr>
		# 		<w:bidi w:val="0"/>
		# 		<w:jc w:val="start"/>
		# 		<w:rPr/>
		# 	</w:pPr>
		# 	<w:r>
		# 		<w:rPr/>
		# 		<w:t>1st level</w:t>
		# 	</w:r>
		# </w:p>

	pPr = create_element('w:pPr')
	numPr = create_element('w:numPr')
	ilvl = create_element('w:ilvl')
	create_attribute(ilvl, 'w:val', str(level))
	numId = create_element('w:numId')
	create_attribute(numId, 'w:val', {'ul':'2', 'ol':'3'}[list_type])
	numPr.append(ilvl)
	numPr.append(numId)
	pPr.append(numPr)
	docx_block._element.append(pPr)
	return docx_block


# ADD HYPERLINK
def add_hyperlink(docx_block, link_text, url):
	# url_base = 'http://127.0.0.1:8080/'
	url_base = 'http://neon/'
	# Omit javascript links
	if not url.startswith('javascript:'):
		# Fix missing domain name
		url = ('http' in url) and url.replace('http:///', url_base) or (url_base + (url.startswith('/') and url[1:] or url))
		r_id = docx_block.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
		hyper_link = create_element('w:hyperlink')
		create_attribute(hyper_link, 'r:id', r_id)
		hyper_link_run = create_element('w:r')
		hyper_link_run_prop = create_element('w:rPr')
		hyper_link_run_prop_style = create_element('w:rStyle')
		create_attribute(hyper_link_run_prop_style, 'w:val', 'Hyperlink')
		hyper_link_run_prop.append(hyper_link_run_prop_style)
		hyper_link_run.append(hyper_link_run_prop)
		hyper_link_text = create_element('w:t')
		hyper_link_text.text = link_text
		hyper_link_run.append(hyper_link_text)
		hyper_link.append(hyper_link_run)

		## DEBUG: Show docx xml
		# etree.tostring(hyper_link, pretty_print=True)

		if isinstance(docx_block, type(docx.Document())):
			# Add hyperlink as a new block element
			p = docx_block.add_paragraph()
			p._p.append(hyper_link)
		else:
			# Add hyperlink as inline element
			docx_block._p.append(hyper_link)
	else:
		docx_block.add_run(link_text)
	return docx_block


# #############################################
# Helper Functions 2: HTML/Richtext-Processing 
# #############################################

# Clean HTML
def clean_html(html):
	# Clean comments, styles, empty tags
	# and handle special characters: left-to-right, triangle
	left_to_right_char = '\u200e'
	triangle_char = '\U0001F806'

	html = re.sub(r'<!.*?->','', html)
	html = re.sub(r'<style.*?</style>','', html)
	html = standard.re_sub(r'\n|\t', ' ', html)
	html = standard.re_sub(r'\s\s', ' ', html)
	html = html.replace('<span class="unicode">&crarr;</span><br />','')
	html = html.replace('<span class="unicode">&crarr;</span>','')
	html = html.replace('">\n','">')
	# refGlossary
	html = html.replace(left_to_right_char,'')
	html = html.replace('[[', triangle_char)
	html = html.replace(']]', '')
	return html

# ADD RUNS TO DOCX-BLOCK
def add_runs(docx_block, bs_element):
	# Adding a minimum set of inline runs
	# any BeautifulSoup block element may contain
	# to the docx-block, e.g. <strong>, <em>, <a>
	if bs_element.children:
		c = 0
		# Hint: ul/ol/li are handled as blocks in add_htmlblock_to_docx.add_list
		elruns = [elrun for elrun in bs_element.children if elrun.name not in ['ul', 'ol', 'li', 'img', 'figure']]
		for elrun in elruns:
			c += 1
			if elrun.name == None:
				s = standard.pystr(elrun)
				# Remove trailing spaces on first text element of a new line or block
				if c == 1:
					s = s.lstrip()
				elif elruns[c-2].name == 'br':
					s = s.lstrip()
				docx_block.add_run(s)
			elif elrun.name == 'br':
				docx_block.add_run('\n')
			elif elrun.name != None and elrun.text == '↵':
				docx_block.add_run('')
			elif elrun.name == 'strong':
				docx_block.add_run(elrun.text).bold = True
			elif elrun.name == 'q':
				docx_block.add_run('"%s"'%(elrun.text), style='Quote-Inline')
			elif elrun.name == 'em' or elrun.name == 'i':
				docx_block.add_run(elrun.text).italic = True
			elif elrun.name in ['samp', 'code', 'tt', 'var', 'pre']:
				docx_block.add_run(elrun.text, style='Code-Inline')
			elif elrun.name == 'kbd':
				docx_block.add_run(elrun.text, style='Keyboard')
				# r.font.highlight_color = WD_COLOR_INDEX.BLACK
			elif elrun.name == 'sub':
				docx_block.add_run(elrun.text).font.subscript = True
			elif elrun.name == 'sup':
				docx_block.add_run(elrun.text).font.superscript = True
			elif elrun.name == 'a':
				if elrun.has_attr('href'):
					add_hyperlink(docx_block = docx_block, link_text = elrun.text, url = elrun.get('href'))
					docx_block.add_run(' ')
				else:
					# elrun.encode_contents() => html
					docx_block.add_run(elrun.text)
			elif elrun.name == 'span':
				if elrun.has_attr('class'):
					class_name = elrun['class'][0]
					style_name = (class_name in doc.styles) and class_name or 'Default Paragraph Font'
					docx_block.add_run(elrun.text, style=style_name)
				else:
					docx_block.add_run(elrun.text)

			# #############################################
			## TO-DO: Add inline image to docx
			## Error: adding image as a block element 
			## may result in content replication
			# #############################################
			# elif elrun.name == 'img':
			# 	add_htmlblock_to_docx(zmscontext, doc, str(elrun), zmsid=None)
			# #############################################
			elif elrun.name == 'p':
				add_runs(docx_block = docx_block, bs_element = elrun)
			# #############################################
	else:
		docx_block.text(standard.pystr(bs_element.text))


# ADD CLASS-NAMED HTML-BLOCK AS PARAGRAPH
def add_tagged_content_as_paragraph(docx_doc, bs_element, style_name="Standard", c=0, zmsid=None):
	# Add content of a BeautifulSoup element as a paragraph
	# to the docx document (self)
	p = docx_doc.add_paragraph(style=style_name)
	if c==1 and zmsid:
		prepend_bookmark(p, zmsid)
	el = BeautifulSoup(standard.pystr(bs_element), 'html.parser')
	el_tag = el.div
	el_tag.unwrap()
	add_runs(docx_block = p,  bs_element = el)


# ADD HTML-BLOCK TO DOCX
def add_htmlblock_to_docx(zmscontext, docx_doc, htmlblock, zmsid=None, zmsmetaid=None):
	# Clean HTML
	htmlblock = clean_html(htmlblock)
	heading_text = ''
	# Apply BeautifulSoup and iterate over elements
	soup = BeautifulSoup(htmlblock, 'html.parser')

	# Counter for html elements: set bookmark before first element
	c = 0

	# Iterate over elements
	for element in soup.children:
		# Skip empty elements
		if element not in ['\n',' ']:
			c+=1
			if isinstance(element, NavigableString):
				# #############################################
				# Text only
				# #############################################
				p = docx_doc.add_paragraph(standard.pystr(element).strip())
				if c==0 and zmsid: 
					prepend_bookmark(p, zmsid)
			else: 
				# #############################################
				# HTML-Elements, element.name != None
				# #############################################

				# #############################################
				# HEADINGS
				# #############################################
				if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
					heading_level = int(element.name[1])
					heading_text = standard.pystr(element.text).strip()
					p = add_heading(docx_doc, heading_text, level=heading_level)
					if c==1 and zmsid: 
						prepend_bookmark(p, zmsid)
					if element.text == 'Inhaltsverzeichnis':
						p.style = docx_doc.styles['TOC-Header']
				# #############################################
				# PARAGRAPH
				# #############################################
				elif element.name == 'p' and element.text != '' and element.text != ' ':
					p = docx_doc.add_paragraph()
					if c==1 and zmsid: 
						prepend_bookmark(p, zmsid)
					# htmlblock.__contains__('ZMSTable') or htmlblock.__contains__('img')
					if element.has_attr('class'):
						if 'caption' in element['class'] and zmsmetaid in ['ZMSGraphic', 'ZMSTable']:
							p.style = docx_doc.styles['caption']
						else:
							class_name = element['class'][0]
							style_name = (class_name in docx_doc.styles) and class_name or 'Normal'
							p.style = docx_doc.styles[style_name]
					add_runs(docx_block = p, bs_element = element)

					## Remove empty paragraphs
					# if element.text == '' or element.text == ' ':
					# 	last_paragraph = docx_doc.paragraphs[-1]._element
					# 	last_paragraph.getparent().remove(last_paragraph)

				# #############################################
				# LIST
				# #############################################
				elif element.name in ['ul','ol']:
					def add_list(docx_doc, element, level=0, c=0):
						for li in element.find_all('li', recursive=False):
							p = docx_doc.add_paragraph()
							set_block_as_listitem(p, list_type=element.name, level=level)
							add_runs(docx_block = p, bs_element = li)
							if c==1 and zmsid:
								prepend_bookmark(p, zmsid)
							for ul in li.find_all(['ul','ol'], recursive=False):
								add_list(docx_doc, ul, level+1)
					add_list(docx_doc, element, level=0, c=c)
				# #############################################
				# TABLE
				# #############################################
				elif element.name == 'table':
					### debug: element.has_attr('class') and element['name']=='abgabekriterium'
					caption = element.find('caption')
					caption_text = caption and caption.text or ''
					p = docx_doc.add_paragraph(standard.pystr(caption_text), style='Table-Caption')
					if c==1 and zmsid:
						prepend_bookmark(p, zmsid)
					rows = element.find_all('tr')
					text_style = 'Normal'
					if len(rows) > 16:
						text_style = 'Table-Small'
					cols = rows[0].find_all(['td','th'])
					# Create table
					docx_table = docx_doc.add_table(rows=len(rows), cols=len(cols))
					docx_table.style = 'Table Grid'
					docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

					def remove_empty_paragraphs(docx_table):
						for row in docx_table.rows:
							for cell in row.cells:
								for p_empty in [p for p in cell.paragraphs if p.text == '']:
									p = p_empty._element
									p.getparent().remove(p)

					# #############################################
					# [A] Filling Cells with data
					# #############################################
					r=-1
					for row in rows:
						r+=1
						cells = row.find_all(['td','th'])
						for i, cl in enumerate(cells):
							docx_cell = docx_table.cell(r,i)
							if {'div','ol','ul','table','p'} & set([e.name for e in cl.children]):
								# [A] Cell contains block elements
								# Hint: Cell implicitly contains a paragraph
								cl_html = ''.join([standard.pystr(child) for child in cl.contents if child!=' '])
								add_htmlblock_to_docx(zmscontext, docx_cell, cl_html, zmsid=None)
							else:
								# [B] Cell contains inline elements
								p = docx_cell.paragraphs[0]
								p.style = docx_doc.styles[text_style]
								if element.has_attr('style') and element['style'] in docx_doc.styles:
									p.style = docx_doc.styles[element['style']]
								add_runs(p, cl)

							# Bolden table header
							if cl.name == 'th':
								docx_table.cell(r,i).paragraphs[0].style = 'Tableheader'

					# #############################################
					# [B] Merge Cells with rowspan and colspan
					# #############################################
					r=-1
					rspn = 0
					clspn = 0
					for row in rows:
						r+=1
						cells = row.find_all(['td','th'])
						for i, cl in enumerate(cells):
							# Merge cells if rowspan or colspan is set
							if cl.has_attr('rowspan'):
								rspn = int(cl['rowspan'])
								docx_table.cell(r,i).merge(docx_table.cell(r,i+rspn-1))
							if cl.has_attr('colspan'):
								clspn = int(cl['colspan'])
								docx_table.cell(r,i).merge(docx_table.cell(r+clspn-1,i))

					# Remove empty paragraphs from any cell
					for row in docx_table.rows:
						for cell in row.cells:
							if len(cell.paragraphs)>1:
								for p_empty in [p for p in cell.paragraphs if p.text == '']:
									p = p_empty._element
									cell._element.remove(p)
									# p.getparent().remove(p)

					# Add linebreak or pagebreak after table
					p = docx_doc.add_paragraph()

				# #############################################
				# IMAGE
				# #############################################
				elif element.name == 'img' or element.name == 'figure':
					if element.name == 'figure':
						element = element.find('img')
					if element.has_attr('src'):
						img_src = element['src']
						try:
							if zmscontext.operator_getattr(zmscontext,zmsid).attr('imghires'):
								# Use high resolution image
								img_src = zmscontext.operator_getattr(zmscontext,zmsid).attr('imghires').getHref(zmscontext.REQUEST)
						except:
							pass
						img_name = img_src.split('/')[-1]
						if not img_src.startswith('http'):
							src_url0 = zmscontext.absolute_url().split('/content/')[0]
							src_url1 = img_src.split('/content/')[-1]
							if src_url1.startswith('/'): 
								# eg. ZMS assets starting with /++resource++zms_
								src_url1 = src_url1[1:]
							element['src'] = '%s/content/%s'%(src_url0, src_url1)

						# Normalize image size to 460px
						imgheight = element.has_attr('height') and int(float(element['height'])) or None
						imgwidth = element.has_attr('width') and int(float(element['width'])) or None
						imgwidth = get_normalized_image_width(w = imgwidth, h = imgheight, max_w = 460)

						try:
							response = requests.get(element['src'], verify=False)
							with open(img_name, 'wb') as f:
								f.write(response.content)
							if src_url1.startswith('++resource++zms_'):
								docx_doc.add_picture(img_name)
							else:
								docx_doc.add_picture(img_name, width=Emu(imgwidth*9525))
							os.remove(img_name)
							if c==1 and zmsid:
								try:
									prepend_bookmark(docx_doc.paragraphs[-1], zmsid)
								except:
									pass
						except:
							pass
				# #############################################
				# DIV
				# #############################################
				elif element.name == 'div':
					if element.has_attr('class') and (('ZMSGraphic' in element['class']) or ('graphic' in element['class'])):
						ZMSGraphic_html = standard.pystr(''.join([str(e) for e in element.children]))
						add_htmlblock_to_docx(zmscontext, docx_doc, ZMSGraphic_html, zmsid, zmsmetaid='ZMSGraphic')
					elif element.has_attr('class') and ('ZMSTextarea' in element['class']):
						ZMSTextarea_html = standard.pystr(''.join([str(e) for e in element.children]))
						add_htmlblock_to_docx(zmscontext, docx_doc, ZMSTextarea_html, zmsid, zmsmetaid='ZMSTextarea')
					elif element.has_attr('class') and 'handlungsaufforderung' in element['class']:
						add_tagged_content_as_paragraph(docx_doc, element, 'Handlungsaufforderung', c, zmsid)
					elif element.has_attr('class') and 'grundsatz' in element['class']:
						add_tagged_content_as_paragraph(docx_doc, element, 'Grundsatz', c, zmsid)
					elif element.has_attr('style') and 'background: rgb(238, 238, 238)' in element['style'] \
						and heading_text != 'Inhaltsverzeichnis':
						add_tagged_content_as_paragraph(docx_doc, element, 'Hinweis', c, zmsid)
					elif element.has_attr('class') and 'text' in element['class'] and zmsmetaid in ['ZMSGraphic', 'ZMSTable']:
						p = docx_doc.add_paragraph(style='Caption')
						if c==1 and zmsid:
							prepend_bookmark(p, zmsid)
						add_runs(docx_block = p, bs_element = element)
					else:
						child_tags = [e.name for e in element.children if e.name]
						if {'em','strong','i'} & set(child_tags):
							p = docx_doc.add_paragraph()
							if c==1 and zmsid: 
								prepend_bookmark(p, zmsid)
							if len(element.contents) == 1:
								if element.has_attr('class'):
									style_name = (class_name in docx_doc.styles) and class_name or 'Normal'
									p.style = docx_doc.styles[style_name]
								p.add_run(element.text)
							elif len(element.contents) > 1:
								for e in element.contents:
									if e.name and e.has_attr('class'):
										class_name = e['class']
										if 'fa-pencil-alt' in class_name:
											# p.add_run('Bearbeitung: ')
											p.add_run(u'\U0000F021', style='Icon')
										elif 'fa-phone' in class_name:
											# p.add_run('Routing: ')
											p.add_run(u'\U0000F028', style='Icon')
											p.add_run('   ')
										if list(e.children)!=[]:
											add_runs(docx_block = p, bs_element = e)
										else:
											p.add_run(standard.pystr(e.text))
									elif e.name:
										p.add_run(standard.pystr(e.text))
									else:
										p.add_run(standard.pystr(e))
							else:
								add_runs(docx_block = p, bs_element = element)
						else:
							div_html = standard.pystr(''.join([standard.pystr(e) for e in element.children]))
							add_htmlblock_to_docx(zmscontext, docx_doc, div_html, zmsid)

				# #############################################
				# Link/A containing text or block elements
				# #############################################
				elif element.name == 'a':
					# Hyperlink just containing text
					if element.children and list(element.children)[0] == element.text:
						try:
							add_hyperlink(docx_block = docx_doc, link_text = element.text, url = element.get('href'))
						except:
							p = docx_doc.add_paragraph()
							if c==1 and	zmsid:
								prepend_bookmark(p, zmsid)
							add_runs(docx_block = p, bs_element = element)
					# Hyperlink containing a block element
					elif {'div','p','table','img'} & set([e.name for e in element.children]):
						# Acquire highres image from fancybox link
						if 'img' in [e.name for e in element.children] and \
							element.has_attr('class') and 'fancybox' in element['class'] and element.has_attr('href'):
							try:
								element.img['src'] = element['href']
							except:
								pass
						div_html = ''.join([str(e) for e in element.children])
						add_htmlblock_to_docx(zmscontext, docx_doc, div_html, zmsid)
					# Hyperlink containing inline elements
					else:
						p = docx_doc.add_paragraph()
						if c==1 and zmsid:
							prepend_bookmark(p, zmsid)
						add_runs(docx_block = p, bs_element = element)
				
				# #############################################
				# FORM
				# #############################################
				elif element.name == 'form':
					p = docx_doc.add_paragraph(style='macro')
					p.add_run('<form>\n').font.bold = True
					if c==1: 
						prepend_bookmark(p, zmsid)
					input_field_count = 0
					for input_field in element.find_all('input', recursive=True):
						input_field_count += 1
						p.add_run('%s. <input> : %s\n'%(input_field_count, input_field.get('name','')))
				# #############################################
				# OTHERS
				# #############################################
				elif element.name == 'hr':
					# Omit horizontal rule
					pass
				elif element.name == 'script':
					# Omit javascript
					pass
				else:
					try:
						if element.has_text:
							p = docx_doc.add_paragraph(standard.pystr(element.text))
							if c==1 and zmsid:
								prepend_bookmark(p, zmsid)
					except:
						docx_doc.add_paragraph(str(element))

	return docx_doc


# ADD BREADCRUMBS AS RUNS TO PARAGRAPH
def add_breadcrumbs_as_runs(zmscontext, p):
	breadcrumbs = zmscontext.breadcrumbs_obj_path()
	c = 0
	for obj in breadcrumbs:
		c += 1
		link_text = obj.meta_id == 'ZMS' and standard.pystr(obj.attr('title')) or standard.pystr(obj.attr('titlealt'))
		add_hyperlink(docx_block = p, link_text = link_text, url = obj.getHref2IndexHtml(zmscontext.REQUEST))
		if c < len(breadcrumbs):
			p.add_run(' > ')
	return p


# #############################################
# Helper Functions 4: GET DOCX NORMALIZED JSON
# #############################################

def apply_standard_json_docx(self):
	"""
	The function creates a normalized JSON stream of 
	a PAGE-like ZMS node. This JSON stream is used for 
	transforming the content to DOCX. 
	It is a list of dicts (key/value-pairs), where the 
	first dict is representing the container meta data
	and the following blocks are representing the PAGEELEMENTS
	of the document.
	Each object dictionary has the following keys:
	- id: the id of the node
	- meta_id: the meta_id of the node
	- parent_id: the id of the parent node
	- parent_meta_id: the meta_id of the parent node
	- title: the title of the node
	- description: the description of the node
	- last_change_dt: the last change date of the node
	- docx_format: the format of the content (html/xml/image 
	  or text-stylename e.g.'Normal')
	- content: the content of the node

	Any PAGEELEMENT-node may have a specific 'standard_json_docx'
	attribute which preprocesses it's ZMS content model close to
	the translation into the DOCX model. The key 'docx_format'
	is used to determine the style of the content block.

	If this attribute method (py-primtive) is not available, 
	the object's class standard_html-method is used to get the 
	content, so that the (maybe not optimum) html will be 
	transformed to DOCX.

	Depending on the complexity of the content it's JSON 
	representation may consist of ore or multiple key/value-
	sequences. Any of these blocks will create a new block 
	element (e.g. paragraph) in the DOCX document.
	"""

	zmscontext = self
	request = zmscontext.REQUEST
	# request.set('preview', 'preview')
	is_page = zmscontext.isPage()

	id = zmscontext.id
	meta_id = zmscontext.meta_id
	parent_id = zmscontext.id
	parent_meta_id = zmscontext.meta_id 
	if zmscontext.meta_id == 'LgRegel':
		title = zmscontext.attr('titlealt')
	else:
		title = zmscontext.attr('title') or zmscontext.getTitle(request)
	description = zmscontext.attr('attr_dc_description')
	last_change_dt = zmscontext.attr('change_dt') or zmscontext.attr('created_dt')
	userid = zmscontext.attr('change_uid')
	url = zmscontext.getHref2IndexHtml(request)

	# Meta data as 1st block
	blocks = [
		{
			'id':id,
			'url':url,
			'meta_id':meta_id,
			'parent_id':parent_id,
			'parent_meta_id':parent_meta_id,
			'title':title,
			'description':description,
			'last_change_dt':last_change_dt,
			'userid':userid,
		}
	]

	if not is_page:
		pageelements = [zmscontext]
	else:
		# Sequence all pageelements including ZMSNote
		# Ref: ZMSObject.isPageElement
		pageelements = [ \
			e for e in zmscontext.getChildNodes(request)  \
				if ( ( e.getType() in [ 'ZMSObject', 'ZMSRecordSet'] ) \
					and not e.meta_id in [ 'LgChangeHistory','ZMSTeaserContainer','LgELearningBanner'] \
					and not e.isPage() ) \
					or e.meta_id in [ 'ZMSLinkElement' ]
			]
		# if zmscontext.meta_id == 'LgRegel':
		# 	pageelements = [zmscontext]

	for pageelement in pageelements:

		json_block = []

		# Check for standard_json_docx attribute
		json_block = pageelement.attr('standard_json_docx')

		if not json_block:

			# #############################################
			# ZMSGraphic
			# #############################################
			if pageelement.meta_id == 'ZMSGraphic':
				id = pageelement.id
				meta_id = pageelement.meta_id
				parent_id = pageelement.getParentNode().id 
				parent_meta_id = pageelement.getParentNode().meta_id 
				text = BeautifulSoup(pageelement.attr('text'), 'html.parser').get_text()
				img = pageelement.attr('imghires') or pageelement.attr('img')
				# img_url = img.getHref(request)
				img_url = '%s/%s'%(pageelement.absolute_url(),img.getHref(request).split('/')[-1])
				imgwidth = img and int(img.getWidth()) or 0
				imgheight = img and int(img.getHeight()) or 0

				json_block = [ 
					{
						'id':'%s_img'%(id), 
						'meta_id':meta_id,
						'parent_id':parent_id,
						'parent_meta_id':parent_meta_id,
						'docx_format':'image',
						'imgwidth':	imgwidth,
						'imgheight':imgheight,
						'content':img_url
					},
					{
						'id':id,
						'meta_id':meta_id,
						'parent_id':parent_id,
						'parent_meta_id':parent_meta_id,
						'docx_format':'Caption',
						'content':'[Abb. %s] %s'%(id, text)
					},

				]
			# #############################################
			# ZMSLinkElement
			# #############################################
			elif pageelement.meta_id == 'ZMSLinkElement':
				# and pageelement.attr('attr_type') in ['','replace','new']:
				id = pageelement.id
				meta_id = pageelement.meta_id
				parent_id = pageelement.getParentNode().id
				parent_meta_id = pageelement.getParentNode().meta_id
				icon = '\U0001F517'
				title = standard.pystr(pageelement.attr('title'))
				text = standard.pystr(pageelement.attr('attr_dc_description'))
				try:
					href = zmscontext.getLinkObj(pageelement.attr('attr_url'),request).getHref2IndexHtml(request)
				except:
					href = '#'

				json_block = [ {
						'id':id,
						'meta_id':meta_id,
						'parent_id':parent_id,
						'parent_meta_id':parent_meta_id,
						'docx_format':'html',
						'content':'<p>%s <a href="%s">%s</a><br/>%s</p>'%(icon, href, title, text)
					}
				]
			# #############################################
			# ZMSFile
			# #############################################
			elif (pageelement.meta_id == 'ZMSFile' or pageelement.meta_id == 'downloadfile') and pageelement.attr('file'):
				id = pageelement.id
				meta_id = pageelement.meta_id
				parent_id = pageelement.getParentNode().id
				parent_meta_id = pageelement.getParentNode().meta_id
				icon = u'\U0001F87E'
				title = standard.pystr(pageelement.attr('title'))
				text = standard.pystr(pageelement.attr('attr_dc_description'))
				try:
					href = pageelement.getHref2IndexHtml(request)
				except:
					href = '#'

				json_block = [ {
						'id':id,
						'meta_id':meta_id,
						'parent_id':parent_id,
						'parent_meta_id':parent_meta_id,
						'docx_format':'html',
						'content':'<p>%s <a href="%s">%s</a>%s</p>'%(icon, href, title, text)
					}
				]

			# #############################################
			# CAVE: Do not apply renderShort!
			# e.g. LgBedingung, ZMSNote
			# #############################################
			elif 'renderShort' in pageelement.getMetaobjAttrIds(pageelement.meta_id):
				json_block = [{
					'id': pageelement.id,
					'meta_id': pageelement.meta_id,
					'parent_id': pageelement.getParentNode().id,
					'parent_meta_id': pageelement.getParentNode().meta_id,
					'docx_format': 'html',
					'content': pageelement.attr('standard_html')
				}]

			# #############################################
			else:
				html = ''
				# Get standard content of pageelement
				try:
					html = pageelement.getBodyContent(request)
					# Clean html data
					html = clean_html(html)
				except:
					html = '<table>'
					html += '<caption>Rendering Error: %s</caption>' % pageelement.meta_id
					attrs = [d['id'] for d in zmscontext.getMetaobjAttrs(pageelement.meta_id) if d['type'] not in ['dtml','zpt','py','constant','resource','interface']]
					for attr in attrs:
						html += '<tr><td>%s</td><td>%s</td></tr>' % (attr, pageelement.attr(attr))
					html += '</table>'
				# Create a json block
				json_block = [{
					'id': pageelement.id,
					'meta_id': pageelement.meta_id,
					'parent_id': pageelement.getParentNode().id,
					'parent_meta_id': pageelement.getParentNode().meta_id,
					'docx_format': 'html',
					'content': html
				}]

			# Give some customizing hints for standard_html
			if pageelement.meta_id in ['LgRegel','LgBedingung','LgELearningBanner','ZMSNote']:
				standard.writeStdout(None, 'IMPORTANT NOTE: %s.standard_html needs to be customized!'%(pageelement.meta_id))
				# %<----  CUSTOMIZE LIKE THIS ---------------------
				# zmi python:request['URL'].find('/manage')>0 and not request['URL'].find('pydocx')>0;
				# %<---- /CUSTOMIZE LIKE THIS ---------------------

		blocks.extend(json_block)

		# Check for newer content
		if pageelement.attr('change_dt') and pageelement.attr('change_dt') >= last_change_dt:
			# Update editorial data
			last_change_dt = pageelement.attr('change_dt')
			if pageelement.attr('change_uid'):
				userid = pageelement.attr('change_uid')

	# Finally set editorial data
	blocks[0]['last_change_dt'] = last_change_dt
	blocks[0]['userid'] = userid

	return blocks

# #############################################
# Helper Functions 5: READ HEADLINE LEVELS
# #############################################
# Based on number-prefix of headlines
def get_headline_levels_from_numbering(headline_paragraphs=[]):
	list2 = [ \
		(
			len(p.text.split(' ')[0].rstrip('.').split('.'))+1 \
			if '.' in p.text.split(' ')[0] else 1 \
		) \
		for p in headline_paragraphs
	]

	return list2


# #############################################
# Helper Functions 6: NORMALIZE HEADLINE LEVELS
# #############################################
def normalize_headline_levels(list1):
	"""
	Normalize headline levels
	expects a list of headline levels as integer values
	"""
	s = []
	list2 = [1]
	for i in list1[1:]:
		i1 = i + 1
		if s and s[-1] == i1:
			pass
		elif not s or s[-1] < i1:
			s.append(i1)
		elif s:
			while len(s) > 1 and s[-1] > i1:
				s = s[:-1]
		list2.append(len(s) + 1)
	return list2

# #############################################
# Helper Functions 7: GET PARAGRAPHS OF SECTION
# #############################################
def get_headings_of_section(doc):
	# Initialize an empty list to hold the sections
	sections = []
	# Initialize an empty list to hold the current section
	section = []

	# Iterate over the elements in the document body
	for element in doc._body._element:
		# If the element is a paragraph
		if element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
			# Create a Paragraph object from the element
			paragraph = Paragraph(element, doc._body)
			# If the paragraph has a heading style
			if 'Heading' in paragraph.style.name:
				# Add the paragraph to the current section
				section.append(paragraph)
		# If the element is a section property
		elif element.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sectPr':
			# Add the current section to the list of sections
			sections.append(section)
			# Start a new section
			section = []

	# Add the last section to the list of sections
	sections.append(section)
	sections[:] = [section for section in sections if section != []]

	# Return the list of sections
	return sections

# Overwrite site-packages/docx/document.py
def add_heading(self, text, level=1):
	"""Return a heading paragraph newly added to the end of the document.

	The heading paragraph will contain `text` and have its paragraph style
	determined by `level`. If `level` is 0, the style is set to `Title`. If `level`
	is 1 (or omitted), `Heading 1` is used. Otherwise the style is set to `Heading
	{level}`. Raises |ValueError| if `level` is outside the range 0-9.
	"""
	if not 0 <= level <= 9:
		raise ValueError("level must be in range 0-9, got %d" % level)
	style = "Title" if level == 0 else "heading %d" % level
	return self.add_paragraph(text, style)


# #############################################
# MAIN function for DOCX-Generation
# #############################################
# The function `manage_export_pydocx` may be called
# recursively to create a DOCX document from a 
# document tree. The function is called with the
# `do_return` parameter set to `True` on the last
# node of the tree. The `filename` parameter is used
# to name the DOCX file. The function returns the
# binary data of the DOCX file.
def manage_export_pydocx(self, save_file=True, file_name=None):
	request = self.REQUEST
	docx_creator = request.AUTHENTICATED_USER.getUserName()

	# PAGE_COUNTER: Counter for recursive export
	page_counter = request.get('page_counter',0) + 1
	request.set('page_counter', page_counter)

	global zmscontext
	zmscontext = self
	is_page = zmscontext.isPage()

	# Write to stdout
	standard.writeStdout(None, 'DOCX-INFO-%s: Writing ZMS node %s' %(page_counter, zmscontext.id))

	# INITIALIZE DOCX Document
	# and preserve it while exporting recursively 
	if page_counter == 1:

		# #############################################
		# DOCX-Document with custom style-set
		# Usage: docx.Document('template.docx')
		# #############################################	
		global doc
		doc = docx.Document(docx_tmpl)
		# Remove eventual example content/paragraphs
		for p in doc.paragraphs:
			e = p._element
			e.getparent().remove(e)
		# set_docx_styles(doc)

		# https://python-docx.readthedocs.io/en/latest/dev/analysis/features/coreprops.html
		doc.core_properties.author = docx_creator
		doc.core_properties.title = standard.pystr(zmscontext.attr('title'))
		doc.core_properties.created = datetime.datetime.now()
		doc.core_properties.modified = datetime.datetime.now()
		doc.core_properties.category = 'ZMS-Export'
		doc.core_properties.comments = 'Generated by ZMS / python-docx'
	
	# GET PAGE CONTENT (JSON)
	zmsdoc = apply_standard_json_docx(self)
	heading = zmsdoc[0]
	blocks = zmsdoc[1:]

	# [A] CREATE SECTION HEADER/FOOTER
	# On recursive export change header for any new page (self)
	# while preserving footer all the same (zmscontext)

	dt = standard.getLangFmtDate(self, heading.get('last_change_dt',''), 'eng', '%Y-%m-%d')
	userid = heading.get('userid','')
	url = heading.get('url','').replace('nohost','localhost')
	url = heading.get('url','').replace('http:///','http://localhost/')
	url = len(url)>124 and url[:124]+'...' or url
	tabs = len(heading.get('title',''))>68 and '\t' or '\t\t'
	header_text = '%s%sOnline Edition %s\nURL: %s'%(standard.pystr(heading.get('title','')), tabs, dt, url)

	if page_counter == 1:
		header_p = doc.sections[0].header.paragraphs[0]
		header_p.clear()
		header_p.text = header_text
		footer_p = doc.sections[0].footer.paragraphs[0]
		footer_p.clear()
		footer_p.add_run('Seite ')
		add_field(footer_p, field_code='PAGE')
		footer_p.add_run('\t')
		footer_p.add_run('Dateiname: ')
		add_field(footer_p, field_code='FILENAME')
		footer_p.add_run(' / Stand: ')
		add_field(footer_p, field_code='SAVEDATE \@ "yyyy-MM-dd" \* MERGEFORMAT')
	else:
		new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
		new_section.header.is_linked_to_previous = False # ESSENTIAL FOR CHANGING HEADER!!!
		header_paragraph = new_section.header.paragraphs[0] if new_section.header.paragraphs else new_section.header.add_paragraph()
		header_paragraph.text = header_text
		# add_breadcrumbs_as_runs(zmscontext, header_paragraph)
		# new_section.header.first_page_header = True
		# new_section.footer.first_page_footer = True

	# [B] CONTENT HEADINGS
	if is_page:
		# [1] CREATE INFO-PAGE
		doc_title = add_heading(doc, standard.pystr(heading.get('title','')), level=0)
		p = doc.add_paragraph()
		add_breadcrumbs_as_runs(zmscontext, p)

		# Document Protocol Table
		v = '''
			<table class="Table-Small">
				<caption>Dokumenthistorie</caption>
				<tr><th>Datum</th><th>Bearbeitungsstadium*</th><th>Bearbeiter</th></tr>
				<tr><td>%s</td><td>ZMS-Edition</td><td>%s</td></tr>
				<tr><td>%s</td><td>ZMS-Export</td><td>%s</td></tr>
				<tr><td> </td><td> </td><td> </td></tr>
				<tr><td> </td><td> </td><td> </td></tr>
				<tr><td> </td><td> </td><td> </td></tr>
				<tr><td> </td><td> </td><td> </td></tr>
			</table>'''%(dt, userid, str(doc.core_properties.created).split(' ')[0], docx_creator)
		add_htmlblock_to_docx(zmscontext=zmscontext, docx_doc=doc, htmlblock=v, zmsid=None)
		p = doc.add_paragraph(style='Table-Small')
		p.add_run(standard.pystr('* mögliche Bearbeitungsstadien siehe Dokument NEON Blueprints: How To')).font.italic = True

		p.add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

		# [2] DOCUMENT TITLE AND DESCRIPTION
		add_heading(doc, standard.pystr(heading.get('title','')), level=1)
		prepend_bookmark(doc.paragraphs[-1], heading.get('id',''))
		if heading.get('description','')!='':
			doc.add_paragraph(standard.pystr(heading.get('description','')), style='Description')

	# [C] CREATE PAGE CONTENT-BLOCKS
	for block in blocks:
		v = standard.pystr(block['content'])
		# #############################################
		# [1] HTML-BLOCK (e.g. richtext with inline styles, just a minimum set of inline elements)
		if v and block['docx_format'] == 'html':
			try:
				add_htmlblock_to_docx(zmscontext=zmscontext, docx_doc=doc, htmlblock=v, zmsid=block['id'])
			except:
				p = doc.add_paragraph()
				p.add_run('Rendering Error: %s'%block['meta_id'])
		# #############################################
		# [2] XML-BLOCK
		elif v and block['docx_format'] == 'xml':
			# Create a paragraph object from parsed xml
			parsed_xml = parse_xml(v)
			doc.add_paragraph()
			# Replace last paragraph with parsed xml
			doc.element.body[-1] = parsed_xml
		# #############################################
		# [3] IMAGE-BLOCK
		elif v and block['docx_format'] == 'image':
			# Add image to document
			image_url = v
			resp = requests.get(image_url, verify=False)
			image_file = BytesIO(resp.content)
			# Add a paragraph containing the image as run
			# see: /python-docx/src/docx/document.py
			p = doc.add_paragraph()
			r = p.add_run()
			# Normalize image width to 460px
			# Get image width and height from 'image' data block
			imgwidth = block.get('imgwidth',0)
			imgheight = block.get('imgheight',0)
			try:
				imgwidth = get_normalized_image_width(w = imgwidth, h = imgheight, max_w = 460)
			except:
				imgwidth = 460
			r.add_picture(image_file, width=Emu(imgwidth*9525))
			prepend_bookmark(p, block['id'])
		# #############################################
		# [4] CAPTION TEXT-BLOCK
		elif v and block['docx_format']=='Caption':
			if re.match(r'^\[Abb. e\d+\] .*', v):
				capt_list = re.split(r'^\[Abb. e\d+\] ', v)
				if len(capt_list) > 1 and len(capt_list[1]) > 0:
					p = doc.add_paragraph(style='Caption')
					prepend_bookmark(p, block['id'])
					p.add_run('Abb. %s: '%block['id']).font.italic = False
					p.add_run(capt_list[1])
			elif re.match(r'^\[Abb. e\d+\] ', v):
				# Omit caption with empty text
				pass
			else:
				p = doc.add_paragraph(style='Caption')
				prepend_bookmark(p, block['id'])
		# #############################################
		# [5] TEXT-BLOCK with given block format (style)
		elif v and block['docx_format'] in [e.name for e in doc.styles]:
				p = doc.add_paragraph(v, style=block['docx_format'])
				prepend_bookmark(p, block['id'])
		elif v:
			p = doc.add_paragraph(v)
			prepend_bookmark(p, block['id'])


	# #############################################
	# Normalize Headline Hierarchy
	# using function normalize_headline_levels 
	# up-levelling systematic gaps in headline levels
	# ---------------------------------------------
	# TODO: Make it work with recursive export
	# #############################################
	if is_page:
		# Get all headline paragraphs for any page / section
		section_list = get_headings_of_section(doc)
		for headline_paragraphs in section_list:
			# 1. Get headline levels on HTML/stylenames 
			# and its normalization with function normalize_headline_levels()
			headline_levels_on_stylenames = [int(p.style.name[-1]) for p in headline_paragraphs]
			headline_levels_normalized = normalize_headline_levels(headline_levels_on_stylenames)

			# 2. Get headline levels on its numbering-prefixes
			headline_levels_on_numbering = get_headline_levels_from_numbering(headline_paragraphs)

			# 3. Which method for getting the normalized levels is more reliable?
			# Sum up prefix-based level-numbers and divide by number of levels
			# Approach: if there are no numbers, all levels are considered as 1
			# what is supposed to be unreliable
			if sum(headline_levels_on_numbering)/len(headline_levels_on_numbering) > 1.25: 
				headline_levels_normalized = headline_levels_on_numbering

			# 4. Reset levels of headline-paragraphs
			for i, p in enumerate(headline_paragraphs):
				p.style = doc.styles['heading %s'%headline_levels_normalized[i]]



	# #############################################
	# [d] SAVE DOCX-FILE
	# #############################################
			
	if save_file:
		# Save document in temporary directory
		fn = '%s.docx'%(file_name and file_name or zmscontext.id_quote(zmscontext.getTitlealt(request)))
		tempfolder = tempfile.mkdtemp()
		docx_file_name = os.path.join(tempfolder, fn)
		doc.save(docx_file_name)
		
		# Read the docx file
		with open(docx_file_name, 'rb') as f:
			docx_file_data = f.read()

		# Remove the temporary folder
		shutil.rmtree(tempfolder)

		# Set the HTTP response headers
		request.RESPONSE.setHeader('Content-Disposition', 'inline;filename=%s'%fn)
		request.RESPONSE.setHeader('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

		# msg = 'DOCX-Export erfolgreich: %s'%fn
		# request.response.redirect(standard.url_append_params('%s/manage_main'%self.absolute_url(),{'lang':request['lang'],'manage_tabs_message':msg}))

		# Return the data of the docx file
		# on single page export or on last page of recursive export
		return docx_file_data
	else:
		# Proceed with recursive export
		pass